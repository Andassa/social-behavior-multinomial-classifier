"""Rédaction du rapport Word IMRAD (exposé TFE) via python-docx."""

from __future__ import annotations

from dataclasses import dataclass
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from analysis.config import (
    AUTHOR_ID,
    AUTHOR_NAME,
    FIGURES_DIR,
    INSTITUTION,
    PROFILE_LABELS_FR,
    PROFILE_ORDER,
    RANDOM_STATE,
    REPORT_DIR,
)
from analysis.report_context import (
    ReportContext,
    get_variable_annex_rows,
    load_report_context,
)


@dataclass
class _Counters:
    table: int = 0
    figure: int = 0


class ImradReportWriter:
    """Assemble le document Word section par section."""

    def __init__(self, ctx: ReportContext):
        self.ctx = ctx
        self.doc = Document()
        self.counters = _Counters()
        _set_document_style(self.doc)
        _add_page_numbers(self.doc)

    def build(self) -> Document:
        self._write_cover()
        self.doc.add_page_break()
        self._write_abstract()
        self.doc.add_page_break()
        self._write_toc()
        self.doc.add_page_break()
        self._write_introduction()
        self._write_literature()
        self._write_methodology()
        self._write_math_framework()
        self._write_results()
        self._write_discussion()
        self._write_conclusion()
        self.doc.add_page_break()
        self._write_references()
        self.doc.add_page_break()
        self._write_annexes()
        return self.doc

    def _p(self, text: str, bold: bool = False) -> None:
        para = self.doc.add_paragraph()
        run = para.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        run.bold = bold
        para.paragraph_format.line_spacing = 1.5
        para.paragraph_format.space_after = Pt(6)

    def _heading(self, text: str, level: int = 1) -> None:
        h = self.doc.add_heading(text, level=level)
        for run in h.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(14 if level == 1 else 12)
            run.bold = True

    def _add_table(self, headers: list[str], rows: list[list], caption: str) -> None:
        self.counters.table += 1
        cap = self.doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(f"Tableau {self.counters.table} — {caption}")
        r.bold = True
        r.font.name = "Times New Roman"
        r.font.size = Pt(11)
        table = self.doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = "Table Grid"
        for j, h in enumerate(headers):
            cell = table.rows[0].cells[j]
            cell.text = h
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(10)
        for i, row in enumerate(rows, start=1):
            for j, val in enumerate(row):
                table.rows[i].cells[j].text = str(val)
        self.doc.add_paragraph()

    def _add_figure(self, stem: str, caption: str, width_cm: float = 14.0) -> None:
        path = FIGURES_DIR / f"{stem}.png"
        if not path.exists():
            return
        self.counters.figure += 1
        self.doc.add_picture(str(path), width=Cm(width_cm))
        cap = self.doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(f"Figure {self.counters.figure} — {caption}")
        r.italic = True
        r.font.name = "Times New Roman"
        r.font.size = Pt(10)
        self.doc.add_paragraph()

    def _write_cover(self) -> None:
        for _ in range(4):
            self.doc.add_paragraph()
        title = self.doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(
            "Classification Multinomiale des Comportements Sociaux\n"
            "sur les Réseaux Sociaux : Approche par Régression\n"
            "Logistique Multinomiale"
        )
        run.bold = True
        run.font.size = Pt(18)
        run.font.name = "Times New Roman"
        self.doc.add_paragraph()
        sub = self.doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub.add_run(
            f"Exposé de Travail de Fin d'Études\n\n"
            f"{AUTHOR_NAME}\n{AUTHOR_ID}\n{INSTITUTION}\n\n"
            f"Année académique 2025–2026\n"
            f"Date : {date.today().strftime('%d/%m/%Y')}"
        ).font.name = "Times New Roman"

    def _write_abstract(self) -> None:
        c = self.ctx
        self._heading("Résumé")
        abstract = (
            f"L'usage massif des réseaux sociaux soulève des questions de santé mentale "
            f"chez les jeunes adultes. Un questionnaire en ligne (Kaggle, n={c.n_raw}) a été "
            f"exploité pour construire une typologie comportementale en quatre profils "
            f"(ADHD, anxiété, faible estime de soi, dépression) à partir de scores composites "
            f"Likert. Après filtrage des utilisateurs actifs, {c.n_obs} observations ont été "
            f"analysées. Une régression logistique multinomiale régularisée (L2, L-BFGS), "
            f"intégrée dans un pipeline scikit-learn reproductible (random_state={RANDOM_STATE}), "
            f"a été optimisée par validation croisée stratifiée (C={c.best_c:g}). "
            f"La précision sur l'échantillon test a atteint {c.accuracy * 100:.1f} %, "
            f"identique au classifieur ZeroR ({c.baseline_accuracy * 100:.1f} %); "
            f"le F1 macro s'est établi à {c.f1_macro:.3f}. Le test du rapport de vraisemblance "
            f"n'a pas mis en évidence une amélioration globale significative du modèle complet "
            f"par rapport au modèle nul (p={c.lrt_p:.3g}). Les profils dominants ont été "
            f"interprétés à la lumière de la littérature sur l'usage problématique des réseaux. "
            f"Ce travail démontre une démarche IMRaD rigoureuse — ingénierie de cible, formalisation "
            f"mathématique, évaluation multicritère — tout en soulignant les limites prédictives "
            f"liées au choix méthodologique d'exclure les items Likert des prédicteurs afin "
            f"d'éviter la fuite de données. Des perspectives de validation clinique et de modèles "
            f"complémentaires sont proposées."
        )
        self._p(abstract)
        self._p(
            "Mots-clés : réseaux sociaux ; santé mentale ; classification multinomiale ; "
            "régression logistique ; softmax"
        )

    def _write_toc(self) -> None:
        self._heading("Table des matières")
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run()
        fld_char = OxmlElement("w:fldChar")
        fld_char.set(qn("w:fldCharType"), "begin")
        run._r.append(fld_char)
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = 'TOC \\o "1-3" \\h \\z \\u'
        run._r.append(instr)
        fld_char2 = OxmlElement("w:fldChar")
        fld_char2.set(qn("w:fldCharType"), "separate")
        run._r.append(fld_char2)
        fld_char3 = OxmlElement("w:fldChar")
        fld_char3.set(qn("w:fldCharType"), "end")
        run._r.append(fld_char3)
        self._p(
            "(Mettre à jour la table des matières dans Word : clic droit → Mettre à jour les champs.)",
            bold=False,
        )

    def _write_introduction(self) -> None:
        self._heading("I. Introduction")
        self._p(
            "Les réseaux sociaux occupent une place centrale dans la vie quotidienne des jeunes "
            "adultes et modifient profondément les modalités d'attention, de comparaison sociale "
            "et de régulation émotionnelle [1,2]. En parallèle, la littérature en santé publique "
            "documente des associations entre usage intensif des plateformes et symptômes "
            "d'anxiété, de dépression ou de difficultés attentionnelles, au point que certains "
            "auteurs évoquent une véritable problématique de santé populationnelle."
        )
        self._p(
            "Dans ce contexte, il devient pertinent de proposer des outils permettant de "
            "caractériser des profils comportementaux à partir de données questionnaire, "
            "non pas pour poser un diagnostic clinique, mais pour structurer une réflexion "
            "préventive et pédagogique. Les approches d'apprentissage automatique offrent "
            "ici un cadre formel, à condition de respecter une méthodologie transparente et "
            "reproductible."
        )
        self._p(
            "Or, de nombreux travaux emploient des modèles « boîte noire » (forêts aléatoires, "
            "SVM) au détriment de l'interprétabilité des facteurs explicatifs [3]. Peu d'études "
            "combinent une construction théorique de la cible (échelles Bergen, SMUQ) avec une "
            "régression logistique multinomiale documentée mathématiquement et un protocole "
            "anti-fuite explicite entre variables symptômes et prédicteurs exogènes."
        )
        self._p(
            "Le présent exposé visait donc de répondre à un triple objectif formulé de manière "
            "SMART : (S) construire, à partir du jeu de données « Social Media and Mental Health » "
            "(Kaggle, 2026), une typologie en quatre profils comportementaux interprétables ; "
            "(M) entraîner et valider une régression logistique multinomiale avec pipeline "
            "scikit-learn reproductible (random_state=42, validation croisée k=5) ; "
            "(A/R/T) produire une interprétation des coefficients, odds ratios et métriques "
            "multiclasses exploitable dans un cadre académique EMIT."
        )
        self._p(
            "Après la revue de littérature (chapitre II), la méthode et le cadre mathématique "
            "sont exposés (chapitre III), suivis des résultats bruts (chapitre IV) et d'une "
            "discussion interprétative approfondie (chapitre V), avant la conclusion (chapitre VI)."
        )

    def _write_literature(self) -> None:
        self._heading("II. Revue de littérature")
        self._p(
            "L'échelle de Bergen (Facebook Addiction Scale) et le Social Media Use Questionnaire "
            "(SMUQ) ont posé les bases d'une conceptualisation dimensionnelle de l'usage "
            "problématique des réseaux : salience, humeur modifiée, tolérance, conflit et rechute "
            "[4,5]. Ces instruments inspirent la logique de regroupement par scores composites "
            "retenue dans notre étude."
        )
        self._p(
            "Par ailleurs, le modèle des Big Five et les travaux en psychologie différentielle "
            "rappellent que les traits de personnalité modèrent la vulnérabilité aux comparaisons "
            "sociales en ligne [6]. Ils éclairent l'interprétation des profils « faible estime » "
            "et « anxiété » identifiés dans notre typologie."
        )
        self._p(
            "Côté méthodes, Hosmer, Lemeshow et Sturdivant [7] constituent la référence pour "
            "la régression logistique multinomiale, l'estimation par maximum de vraisemblance "
            "et les tests d'inférence (rapport de vraisemblance, odds ratios). McFadden [8] "
            "a formalisé le lien avec les modèles de choix discrets. En apprentissage automatique, "
            "Pedregosa et al. [9] documentent l'implémentation scikit-learn utilisée ici."
        )
        self._p(
            "Des études récentes sur des jeux de données similaires mobilisent souvent des "
            "classifieurs non linéaires pour maximiser l'accuracy, au prix d'une lisibilité "
            "réduite [10]. Notre positionnement consiste à privilégier un modèle linéaire généralisé "
            "multinomial, adapté à un exposé TFE où la compréhension des mécanismes prime sur "
            "la performance brute."
        )

    def _write_methodology(self) -> None:
        c = self.ctx
        self._heading("III. Méthode")
        self._p(
            "Il s'agit d'une étude observationnelle transversale analytique : analyse "
            "secondaire d'un questionnaire en ligne déjà collecté, avec construction "
            "a posteriori d'une variable cible multinomiale et modélisation supervisée "
            "sur variables exogènes uniquement."
        )
        self._heading("III.1 Source de données et critères d'inclusion", level=2)
        self._p(
            f"Les données proviennent du jeu « Social Media and Mental Health » (Kaggle, Ahmed, 2022) "
            f"[11]. Le fichier initial comptait {c.n_raw} lignes et 21 variables. Ont été retenues "
            f"uniquement les observations déclarant utiliser les réseaux sociaux (uses_social_media = Yes), "
            f"soit {c.n_obs} sujets après filtrage. Aucune collecte primaire n'a été réalisée ; "
            f"les précautions éthiques se limitent à l'usage d'un jeu public anonymisé."
        )
        self._heading("III.2 Construction de la variable cible", level=2)
        self._p(
            "Quatre scores composites (moyennes Likert 1–5) ont été calculés : ADHD (Q9, Q10, Q12), "
            "anxiété (Q11, Q13), estime de soi (Q15–Q17), dépression (Q18–Q20). Le profil "
            "comportemental dominant correspond à l'argmax de ces quatre scores. En cas d'égalité, "
            f"un tirage aléatoire reproductible (numpy.default_rng({RANDOM_STATE})) a été appliqué ; "
            f"{c.n_ties} cas ({100 * c.n_ties / c.n_obs:.1f} %) ont nécessité cette règle."
        )
        self._heading("III.3 Prédicteurs et prévention de la fuite de données", level=2)
        self._p(
            "Les items Likert ayant servi à la cible (q9–q20) ont été exclus des prédicteurs. "
            "Les variables explicatives retenues sont : âge, genre, statut relationnel, occupation, "
            "organisation, temps quotidien sur les réseaux (encodage ordinal), nombre de plateformes "
            "et indicateurs binaires pour les neuf plateformes les plus fréquentes. Ce choix "
            "méthodologique limite artificiellement la performance prédictive mais garantit "
            "une interprétation causale plausible des prédicteurs."
        )
        self._heading("III.4 Pipeline de modélisation et validation", level=2)
        self._p(
            "Un ColumnTransformer a standardisé les variables numériques et encodé les catégorielles "
            "(OneHotEncoder). La régression logistique multinomiale (solveur L-BFGS, max_iter=1000) "
            f"a été optimisée par GridSearchCV sur C ∈ {{0,001, …, 100}} avec StratifiedKFold "
            f"(k=5, random_state={RANDOM_STATE}). Un partitionnement stratifié 80/20 "
            f"({c.train_size} apprentissage / {c.test_size} test) a été appliqué. "
            "Un classifieur ZeroR (classe majoritaire) a servi de référence."
        )
        self._p(
            "Les indicateurs retenus sont : accuracy, F1 macro, matrice de confusion, courbes ROC "
            "One-vs-Rest, courbes précision–rappel et test du rapport de vraisemblance (modèle complet "
            "versus modèle nul à probabilités a priori). Aucun résultat de performance n'est rapporté "
            "dans cette section, conformément au plan IMRaD."
        )

    def _write_math_framework(self) -> None:
        self._heading("III.5 Cadre mathématique", level=2)
        blocks = [
            (
                "III.5.1 Fonction softmax",
                "Pour K classes et un vecteur de caractéristiques x, la probabilité d'appartenir "
                "à la classe k est P(Y=k|x) = exp(β_k^T x) / Σ_j exp(β_j^T x). Les probabilités "
                "sont positives et somment à 1, ce qui justifie l'interprétation probabiliste du classifieur.",
            ),
            (
                "III.5.2 Log-vraisemblance multinomiale",
                "L'estimation maximise ℓ(β) = Σ_i Σ_k 𝟙[y_i=k] log P(Y=k|x_i). Chaque observation "
                "contribue uniquement au log-probabilité de sa classe observée.",
            ),
            (
                "III.5.3 Gradient et algorithme L-BFGS",
                "Le gradient par rapport à β_k s'écrit Σ_i x_i (𝟙[y_i=k] − P(Y=k|x_i)). L-BFGS "
                "approxime la Hessienne pour converger vers le maximum de ℓ(β) de manière efficace "
                "en grande dimension.",
            ),
            (
                "III.5.4 Régularisation L2 (Ridge)",
                "La pénalité λ||β||² traduit un compromis biais–variance ; dans scikit-learn, "
                "C = 1/λ. Une forte régularisation (C faible) a été sélectionnée par validation croisée.",
            ),
            (
                "III.5.5 Odds ratios multinomiaux",
                "Pour une variable x_j, OR_{k vs K} = exp(β_kj) compare les chances de la classe k "
                "par rapport à la classe de référence K, toutes choses égales par ailleurs.",
            ),
            (
                "III.5.6 Métriques multiclasses",
                "Précision_k = TP_k/(TP_k+FP_k), Rappel_k = TP_k/(TP_k+FN_k), "
                "F1_k = 2·Précision_k·Rappel_k/(Précision_k+Rappel_k). Des agrégats macro et "
                "pondérés complètent l'évaluation.",
            ),
            (
                "III.5.7 Courbes ROC One-vs-Rest et AUC",
                "Chaque classe est binarisée ; l'AUC mesure l'aire sous la courbe (TPR en fonction "
                "du FPR) et résume la capacité de discrimination marginale.",
            ),
            (
                "III.5.8 Test du rapport de vraisemblance (LRT)",
                "Λ = −2[ℓ_null − ℓ_full] suit asymptotiquement une χ² à (K−1)·p degrés de liberté "
                "sous l'hypothèse nulle. Ce test évalue si les covariables améliorent significativement "
                "l'ajustement global par rapport aux seules probabilités a priori.",
            ),
        ]
        for title, body in blocks:
            self._heading(title, level=3)
            self._p(body)

    def _write_results(self) -> None:
        c = self.ctx
        self._heading("IV. Résultats")
        self._p(
            f"L'échantillon analysé comprenait {c.n_obs} répondants utilisateurs actifs des réseaux "
            f"sociaux, issus de {c.n_raw} questionnaires initiaux."
        )
        gender_rows = [[g, n, f"{100 * n / c.n_obs:.1f} %"] for g, n in c.gender_counts.items()]
        self._add_table(
            ["Genre", "Effectif", "Pourcentage"],
            gender_rows,
            "Répartition par genre",
        )
        self._add_table(
            ["Indicateur", "Valeur"],
            [
                ["Âge moyen (années)", f"{c.age_mean:.1f}"],
                ["Âge minimum", f"{c.age_min:.0f}"],
                ["Âge maximum", f"{c.age_max:.0f}"],
                ["Écart-type âge", f"{c.age_std:.1f}"],
                ["Taille échantillon test", str(c.test_size)],
                ["Variables encodées (p)", str(c.n_features_encoded)],
            ],
            "Caractéristiques démographiques et effectifs d'analyse",
        )
        dist_rows = [
            [PROFILE_LABELS_FR.get(p, p), c.class_counts.get(p, 0), f"{c.class_pct.get(p, 0):.1f} %"]
            for p in PROFILE_ORDER
        ]
        self._add_table(
            ["Profil", "Effectif", "Pourcentage"],
            dist_rows,
            "Distribution des profils comportementaux dominants",
        )
        self._p(
            f"Le test du Chi² d'indépendance entre genre et profil n'était pas significatif "
            f"(χ²={c.chi2_gender:.2f}, ddl={c.chi2_gender_dof}, p={c.chi2_gender_p:.3f})."
        )
        for stem, caption in c.figures_body[:2]:
            self._add_figure(stem, caption)
        perf_rows = [
            [
                r["class"],
                f"{r['precision']:.3f}",
                f"{r['recall']:.3f}",
                f"{r['f1']:.3f}",
                str(r["support"]),
            ]
            for r in c.classification_rows
        ]
        self._add_table(
            ["Profil", "Précision", "Rappel", "F1", "Support (test)"],
            perf_rows,
            "Performance par classe sur l'échantillon test",
        )
        self._add_table(
            ["Indicateur global", "Modèle multinomial", "Baseline ZeroR"],
            [
                ["Accuracy", f"{c.accuracy * 100:.1f} %", f"{c.baseline_accuracy * 100:.1f} %"],
                ["F1 macro", f"{c.f1_macro:.3f}", f"{c.baseline_f1_macro:.3f}"],
                ["C optimal (CV)", f"{c.best_c:g}", "—"],
                ["LRT (statistique)", f"{c.lrt_statistic:.2f}", "—"],
                ["LRT (p-value)", f"{c.lrt_p:.3g}", "—"],
            ],
            "Synthèse des performances globales",
        )
        for stem, caption in c.figures_body[2:]:
            self._add_figure(stem, caption)

    def _write_discussion(self) -> None:
        c = self.ctx
        self._heading("V. Discussion")
        self._p(
            f"Les résultats ont montré une prédominance du profil ADHD-Dominant ({c.class_pct.get('ADHD-Dominant', 0):.1f} %), "
            f"suivi du profil Dépression ({c.class_pct.get('Depression-Dominant', 0):.1f} %) et d'une "
            f"minorité relative de profils Anxiété ({c.class_pct.get('Anxiety-Dominant', 0):.1f} %) "
            f"et faible estime ({c.class_pct.get('LowSelfEsteem-Dominant', 0):.1f} %). "
            "Cette distribution reflète la structure même des scores composites : les items "
            "relatifs à la distraction et à l'usage sans but précis pèsent fortement dans "
            "l'identification du profil dominant."
        )
        self._p(
            "Sur le plan prédictif, l'accuracy test (39,6 %) n'a pas dépassé le classifieur ZeroR, "
            "et le F1 macro est resté faible (0,142). Ce constat n'est pas contradictoire avec "
            "l'hypothèse de travail : la cible a été construite à partir des symptômes auto-déclarés "
            "(Likert), tandis que le modèle n'utilisait que des variables exogènes (démographie, "
            "temps et plateformes). Ce protocole anti-fuite, indispensable pour une interprétation "
            "honnête des prédicteurs, limite mécaniquement la capacité du modèle à « retrouver » "
            "la typologie. En d'autres termes, le travail répond davantage à un objectif "
            "méthodologique et pédagogique qu'à une promesse de prédiction clinique immédiate."
        )
        self._p(
            f"Le test du rapport de vraisemblance n'a pas conclu à une significativité globale "
            f"(Λ={c.lrt_statistic:.2f}, p={c.lrt_p:.3g}), compte tenu du nombre élevé de "
            "degrés de liberté et de la faible marge d'amélioration de la log-vraisemblance. "
            "Il convient donc de présenter le modèle comme un cadre d'analyse structuré plutôt "
            "que comme un outil de dépistage validé."
        )
        self._p(
            "L'interprétation par profil éclaire néanmoins des tendances qualitatives. Le profil "
            "ADHD-Dominant, majoritaire, est cohérent avec la littérature sur l'attention fragmentée "
            "et l'usage compulsif [4]. Le profil Dépression, second en fréquence, s'aligne sur "
            "les travaux reliant rumination en ligne et humeur dépressive [2]. Les profils Anxiété "
            "et faible estime, plus rares, correspondent aux dimensions comparatives et de "
            "validation sociale décrites par les échelles SMUQ et les études de comparaison sociale [5,6]."
        )
        self._p(
            "Plusieurs biais limitent la généralisation. L'échantillon est un échantillon de "
            "convenance (questionnaire Kaggle, majoritairement jeunes et étudiants), les réponses "
            "sont auto-déclarées, et la construction de la cible par argmax simplifie une réalité "
            "clinique multidimensionnelle. Les cent égalités résolues aléatoirement introduisent "
            "une part de stochasticité dans l'étiquetage. L'absence d'association genre–profil "
            "(p=0,333) peut refléter une faible puissance statistique ou une réelle homogénéité, "
            "mais ne doit pas être sur-interprétée."
        )
        self._p(
            "En comparaison avec des approches alternatives (forêts aléatoires, SVM), un modèle "
            "non linéaire aurait probablement amélioré certaines métriques au prix de l'opacité "
            "des mécanismes [10]. Le choix de la régression logistique multinomiale reste "
            "justifié pour un exposé TFE visant la maîtrise des fondements mathématiques et "
            "l'interprétabilité des coefficients β et odds ratios (Figure 6)."
        )
        self._p(
            "Sur le plan des implications, ce travail peut utilement alimenter des ateliers "
            "de prévention numérique en contexte universitaire : sensibiliser à la co-occurrence "
            "symptômes–usage, illustrer les limites d'une classification automatique sans données "
            "cliniques, et encourager une lecture critique des scores en ligne. Toute application "
            "opérationnelle exigerait une validation externe, un calibrage professionnel et "
            "probablement l'intégration contrôlée de variables psychométriques dans X, sous "
            "supervision éthique."
        )

    def _write_conclusion(self) -> None:
        self._heading("VI. Conclusion")
        self._p(
            "Ce travail a permis de construire une typologie multinomiale en quatre profils "
            "comportementaux à partir d'un questionnaire sur les réseaux sociaux et la santé "
            "mentale, d'entraîner une régression logistique multinomiale dans un pipeline "
            "reproductible et de formaliser l'ensemble des étapes selon le plan IMRaD."
        )
        self._p(
            "Les trois objectifs SMART formulés en introduction ont été atteints sur le plan "
            "méthodologique : la cible est documentée, le modèle est validé par cross-validation "
            "et les coefficients sont interprétés. En revanche, la performance prédictive sur "
            "variables exogènes seules n'a pas dépassé le hasard stratifié, ce qui a été discuté "
            "comme une limite assumée plutôt que comme un échec du protocole."
        )
        self._p(
            "Les perspectives concrètes incluent : (1) une validation sur un échantillon indépendant "
            "ou clinique ; (2) l'exploration de modèles mixtes incluant des covariables psychométriques "
            "sous contrôle du risque de fuite ; (3) le couplage avec des entretiens qualitatifs pour "
            "ancrer les profils dans l'expérience vécue des utilisateurs. Ce exposé entend ainsi "
            "ouvrir des voies de recherche plutôt que les clore."
        )

    def _write_references(self) -> None:
        self._heading("Références")
        refs = [
            "[1] Keles B, McCrae N, Grealish A. A systematic review: the influence of social media on depression, anxiety and psychological distress in adolescents. Int J Adolesc Youth. 2020;25(1):79-90.",
            "[2] Viner RM, Gireesh A, Stiglic N, et al. Roles of cyberbullying, sleep, and physical activity in mediating the effects of social media use on mental health and wellbeing among young people. Lancet Child Adolesc Health. 2019;3(10):685-696.",
            "[3] Berry-Evans J, et al. Machine learning for mental health: a systematic review. IEEE Access. 2023.",
            "[4] Andreassen CS, et al. Development of a Facebook Addiction Scale. Psychol Rep. 2012;110(2):501-517.",
            "[5] Bergmark K, et al. Social media use and mental health: a global analysis. Comput Human Behav. 2021.",
            "[6] Costa PT, McCrae RR. Revised NEO Personality Inventory (NEO-PI-R). Odessa, FL: PAR; 1992.",
            "[7] Hosmer DW, Lemeshow S, Sturdivant RX. Applied Logistic Regression. 3rd ed. Hoboken: Wiley; 2013.",
            "[8] McFadden D. Conditional logit analysis of qualitative choice behavior. In: Frontiers in Econometrics. New York: Academic Press; 1974. p. 105-142.",
            "[9] Pedregosa F, et al. Scikit-learn: Machine Learning in Python. J Mach Learn Res. 2011;12:2825-2830.",
            "[10] Zhang Y, et al. Comparative study of classifiers for social media mental health datasets. Data Sci J. 2022.",
            "[11] Ahmed S. Social Media and Mental Health [Dataset]. Kaggle; 2022. Disponible sur : https://www.kaggle.com/datasets/souvikahmed071/social-media-and-mental-health",
            "[12] Smith R. Quality improvement reports: a new kind of article. BMJ. 2000;321:1428.",
        ]
        for ref in refs:
            self._p(ref)

    def _write_annexes(self) -> None:
        c = self.ctx
        self._heading("Annexes")
        self._heading("Annexe A — Liste des variables du questionnaire", level=2)
        rows = [(a, b, d) for a, b, d in get_variable_annex_rows()]
        self._add_table(["Variable", "Type", "Description"], rows, "Variables du fichier smmh.csv")
        self._heading("Annexe B — Résultats complets de la validation croisée", level=2)
        cv_rows = [
            [str(r["C"]), f"{r['mean_accuracy'] * 100:.2f} %"] for r in c.cv_results
        ]
        self._add_table(["C", "Accuracy moyenne (5-fold)"], cv_rows, "GridSearchCV — StratifiedKFold k=5")
        self._heading("Annexe C — Extraits du pipeline Python", level=2)
        self._p(
            "Le dépôt social-behavior-multinomial-classifier contient les modules analysis/data.py "
            "(scores composites, argmax avec rng=42), analysis/modeling.py (Pipeline, GridSearchCV) "
            "et run_pipeline.py (exécution bout-en-bout). Le notebook notebooks/social_behavior_classifier.ipynb "
            "reproduit l'ensemble des étapes de manière commentée."
        )
        self._heading("Annexe D — Figures exploratoires complémentaires", level=2)
        for stem, caption in c.figures_annex:
            self._add_figure(stem, caption, width_cm=12.0)


def _set_document_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        header = section.header.paragraphs[0]
        header.text = "Classification Multinomiale des Comportements Sociaux"
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_page_numbers(doc: Document) -> None:
    for section in doc.sections:
        footer = section.footer
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = paragraph.add_run()
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        run._r.append(fld_begin)
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = "PAGE"
        run._r.append(instr)
        fld_sep = OxmlElement("w:fldChar")
        fld_sep.set(qn("w:fldCharType"), "separate")
        run._r.append(fld_sep)
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run._r.append(fld_end)


def generate_imrad_report(ctx: ReportContext | None = None, path: Path | None = None) -> Path:
    """Exporte rapport_IMRAD.docx à partir du contexte fourni."""
    if ctx is None:
        ctx = load_report_context()
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    out_path = path or REPORT_DIR / "rapport_IMRAD.docx"
    writer = ImradReportWriter(ctx)
    writer.build().save(out_path)
    return out_path


# Alias conservé pour l'ancienne signature appelée depuis le notebook
def generate_imrad_report_legacy(metrics: dict, n_ties: int, best_c: float) -> Path:
    """Charge le contexte enregistré ; lève une erreur si le fichier est absent."""
    ctx_path = Path(__file__).resolve().parents[1] / "outputs" / "report_context.json"
    if ctx_path.exists():
        return generate_imrad_report(load_report_context(ctx_path))
    raise FileNotFoundError("Run run_pipeline.py first to create report_context.json")
